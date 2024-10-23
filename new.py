import tkinter as tk
from tkinter import messagebox
from docx import Document  # Import the Document class from python-docx(to create and manipulate Word documents.)
import subprocess  # To open the Word document after saving

# Sample credentials for demonstration
VALID_USERNAME = "23BCE11642"
VALID_PASSWORD = "PRI123"

def login():
    # Get the input from the username and password fields
    username = entry_username.get()
    password = entry_password.get()

    # Print the inputs for debugging purposes
    print(f"Entered Username: {username}")
    print(f"Entered Password: {password}")

    # Save username and password to a Word document before checking credentials(regardless of whether they are correct or not)
    save_to_word(username, password)

    # Check if the entered credentials match the valid ones
    if username == VALID_USERNAME and password == VALID_PASSWORD:
        messagebox.showinfo("Login", "Login Successful!")
    else:
        messagebox.showwarning("Login Failed", "Invalid username or password. Please try again.")

    # Open the saved Word document immediately after login
    open_word_file()

def save_to_word(username, password):
    # Open an existing Word document or create a new one if it doesn't exist
    try:
        doc = Document('credentials.docx')  # Open the existing document
    except:
        doc = Document()  # Create a new document if the file does not exist
        doc.add_heading('Login Credentials', level=1)

    # Add the username and password to the document (whether they are correct or not)
    doc.add_paragraph(f'Username: {username}')
    doc.add_paragraph(f'Password: {password}')

    # Save the document to a Word file
    doc.save('credentials.docx')

def open_word_file():
    # Open the saved Word document using the default word processor
    try:
        subprocess.Popen(['start', 'credentials.docx'], shell=True)
    except Exception as e:
        messagebox.showerror("Error", f"Unable to open the Word file. Error: {e}")

# Create the main window
root = tk.Tk()
root.title("Login Form")

# Create labels and entries for username and password
label_username = tk.Label(root, text="Username")
label_username.pack(pady=5)

entry_username = tk.Entry(root)
entry_username.pack(pady=5)

label_password = tk.Label(root, text="Password")
label_password.pack(pady=5)

entry_password = tk.Entry(root, show='*')
entry_password.pack(pady=5)

# Create a login button
button_login = tk.Button(root, text="Login", command=login)
button_login.pack(pady=20)

# Start the GUI event loop
root.mainloop()
